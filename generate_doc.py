"""生成智能体说明文档 — 首届超级智能体创新大赛"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── Page setup ──
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def add_para(text, bold=False, size=12, align='left', font_name='宋体',
             space_after=6, first_line_indent=None):
    p = doc.add_paragraph()
    p.alignment = {'left': 0, 'center': 1, 'right': 2}.get(align, 0)
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.25
    if first_line_indent:
        pf.first_line_indent = Cm(first_line_indent)
    run = p.add_run(text)
    run.font.name = font_name
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.bold = bold
    return p


def add_heading_styled(text):
    p = doc.add_paragraph()
    p.alignment = 0
    pf = p.paragraph_format
    pf.space_before = Pt(8)
    pf.space_after = Pt(4)
    pf.line_spacing = 1.25
    run = p.add_run(text)
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(12)
    run.bold = True
    return p


# ============================================================
# COVER PAGE
# ============================================================

add_para('附件3', size=14, align='center', space_after=12)
add_para('首届超级智能体创新大赛', bold=True, size=18, align='center', space_after=4)
add_para('智能体说明文档', bold=True, size=16, align='center', space_after=20)

# Info table
table = doc.add_table(rows=5, cols=2)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

for row in table.rows:
    row.cells[0].width = Cm(3.0)
    row.cells[1].width = Cm(8.0)

info_data = [
    ('智能体名称', '智游伴（TravelMind Agent）'),
    ('访问地址', '本地开发中（Docker 一键部署，公网待发布）'),
    ('负 责 人', '王文治'),
    ('成    员', '王文治'),
    ('填表日期', '2026-07-25'),
]

for i, (label, value) in enumerate(info_data):
    cell_label = table.rows[i].cells[0]
    cell_value = table.rows[i].cells[1]

    p = cell_label.paragraphs[0]
    p.alignment = 1
    run = p.add_run(label)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    run.bold = True

    p = cell_value.paragraphs[0]
    p.alignment = 1
    run = p.add_run(value)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)

add_para('', space_after=16)
add_para('广东省高等教育学会教育技术专业委员会 制表', size=10, align='center', space_after=4)
add_para('2026年5月', size=10, align='center', space_after=20)

# ── Page Break ──
doc.add_page_break()

# ============================================================
# MAIN CONTENT
# ============================================================

# ── 一、作品概述 ──
add_heading_styled('一、作品概述')
add_para(
    '智游伴（TravelMind Agent）是一款基于大语言模型、多 Agent 协作、RAG 知识增强与多模态理解'
    '能力的 AI 旅行规划智能体。用户以自然语言输入旅行需求（如"带父母去重庆玩3天，喜欢美食和夜景"），'
    '系统通过七步异步智能管线自动完成画像提取、趋势分析、天气查询、知识检索、智能推荐、行程生成与'
    '效果校验，最终输出一份经过 POI 存续验证、路线优化、真实价格注入的结构化行程卡片。系统支持对话式'
    '规划、图片灵感转行程、智能推荐、行程版本管理等多种交互方式，知识库覆盖 30 个国内热门旅游城市、'
    '1,788 个 POI（含 454 美食 POI），评测体系含 63 条查询 × 24 项约束，当前基线 Micro 96.4%、'
    'Macro 74.6%。',
    first_line_indent=0.74
)

# ── 二、痛点问题与应用场景 ──
add_heading_styled('二、痛点问题与应用场景')
add_para('（一）痛点问题', bold=True, size=11, space_after=2)
pain_points = [
    '传统旅游规划耗时 6–8 小时，需跨平台查阅攻略、比对景点、手动编排路线；',
    '网络攻略信息过时且不可靠——景点关闭或搬迁、门票价格变动频繁；',
    '推荐算法千篇一律，缺乏场景自适应能力（如雨天应优先推荐室内景点）；',
    '用户看到喜欢的旅行照片，缺乏高效工具找到相似体验的目的地。',
]
for pp in pain_points:
    add_para('• ' + pp, size=10.5, space_after=1, first_line_indent=0.74)

add_para('（二）应用场景', bold=True, size=11, space_after=2)
scenes = [
    '对话式旅行规划：多轮对话收敛意图 → 一键生成结构化行程（面向普通用户）；',
    '智能推荐：输入偏好标签 → 7 因子评分排序景点，支持跨城搜索（面向有明确偏好的用户）；',
    '图片灵感转行程：上传风景照 → Kimi k2.6 识别地标/风格 → 全库跨城推荐相似景点（面向被"种草"的用户）；',
    '行程管理：支持版本历史快照、一键恢复、收藏、局部重生成（面向反复打磨的用户）。',
]
for s in scenes:
    add_para('• ' + s, size=10.5, space_after=1, first_line_indent=0.74)

# ── 三、智能体设计 ──
add_heading_styled('三、智能体设计')
add_para(
    '本智能体采用 9 Agent 协作 + 对话状态机的混合架构。核心 Agent 设计如下：',
    first_line_indent=0.74, size=10.5
)

agents = [
    ('Profile Agent',
     'NL → 结构化画像（目的地/预算/天数/标签/同伴），DeepSeek v4-flash structured output。'),
    ('Trend Agent',
     '118+ 条趋势数据 + 4 策略模糊匹配，补充热门景点并加权。'),
    ('Weather Agent',
     'Open-Meteo 7 天预报 + 旅行适宜度评分（雷暴/大雨/高温/大风扣分模型）。'),
    ('RAG Retriever',
     'ChromaDB 向量检索（1075 维混合向量：TF-IDF + Tag One-Hot）+ 6 因子精排，'
     '首创天气感知 Boost——降雨日室内 POI +0.20×rain_ratio、户外 -0.10×rain_ratio。'),
    ('Recommendation Agent',
     '7 因子加权打分（偏好/热度/预算/位置/季节/可靠性/天气），过滤无效 POI。'),
    ('Planning Agent',
     'DeepSeek 生成结构化行程（JSON Schema 约束），最多 3 次重试。'),
    ('Route Optimizer',
     'POI 存续校验 + 闭店替换 + 区域归位 + 近邻重排。'),
    ('Dialog Manager',
     '五阶段对话状态机（COLLECTING→CONFIRMING→GENERATING→DELIVERED），'
     '支持增量修改分流（local/slot_change/global）。'),
    ('Vision Agent',
     'Kimi k2.6 识别地标/风格标签 → 闭环推荐。'),
]
for name, desc in agents:
    add_para('• ' + name + '：' + desc, size=10, space_after=1, first_line_indent=0.74)

add_para(
    '核心设计理念：① 天气感知双 Boost——在检索和推荐两层同时施加天气自适应评分，'
    '使推荐结果自适应实时天气；② 数据完整性铁律——严禁 AI 编造任何数据，所有数值均来自'
    '真实 API（高德、Open-Meteo、Wikidata）；③ 增量修改——行程生成后支持局部重生成，'
    '避免全网重跑管线。',
    size=10.5, first_line_indent=0.74
)

# ── 四、核心功能与技术架构 ──
add_heading_styled('四、核心功能与技术架构')
add_para(
    '系统采用三层技术架构：L1 交互层——React 19 + TypeScript + Vite + Tailwind CSS（6 页面 / '
    '14 组件 + SSE 流式渲染 + PDF 导出）；L2 意图层——FastAPI + 23 条 API 路由 + 对话状态机'
    '（Pydantic v2 数据校验 + Alembic 数据库迁移）；L3 生成层——7 步异步管线 + 行程契约校验'
    ' + 真实价格注入 + POI 存续巡检。',
    size=10.5, first_line_indent=0.74
)

add_para(
    '数据基础设施：知识库含 1,788 POI / 30 城市，经六步管线构建（Wikidata SPARQL → Wikipedia '
    '摘要 → 高德 Amap POI 补充 → DeepSeek AI 标注 → 确定性价格注入 → Chroma 向量化）。嵌入方案'
    '采用纯 Python TF-IDF（char n-gram 2-4, 1024 维）+ Tag One-Hot（51 维）= 1075 维混合向量，'
    '零 GPU 依赖、零外部 Embedding API 成本。外部 API 依赖：DeepSeek（LLM）+ Kimi（视觉）+ '
    '高德 Amap（地图）+ Open-Meteo（天气，免费无限）。DevOps：Docker 四服务编排（backend + '
    'frontend + redis + postgres）+ GitHub Actions CI（pytest + tsc + oxlint），支持一键部署。',
    size=10.5, first_line_indent=0.74
)

# ── 五、应用效果与创新点 ──
add_heading_styled('五、应用效果与创新点')
add_para('（一）应用效果', bold=True, size=11, space_after=2)

effects = [
    '评测体系：63 条查询 × 24 项约束，全部由确定性代码打分（不用 LLM 当评委），结果可复现；',
    '核心指标：Micro 通过率 96.4%、Macro 通过率 74.6%（47/63 全通过）；',
    '关键约束：天气匹配率 75.0%（Phase 12.16 天气感知双 Boost 提升 +7.5pp），'
    'POI 存续验证率 97.5%，路线合理性 100%；',
    '工程质量：301 个单元测试（0 失败）、TypeScript 0 编译错误、Docker 一键部署。',
]
for e in effects:
    add_para('• ' + e, size=10.5, space_after=1, first_line_indent=0.74)

add_para('（二）创新点', bold=True, size=11, space_after=2)
innovations = [
    '天气感知 RAG + 推荐双 Boost：首创在检索和推荐两层同时施加天气自适应评分，rain_ratio '
    '精准控制调节强度，实验证明 weather_fit 指标提升 +7.5pp（67.5% → 75.0%）；',
    '六步数据管线：Wikidata → Wikipedia → 高德 → AI → 价格 → Chroma，全部合法 API/CC 授权源，'
    '零爬取，数据可溯源；',
    '对话状态机 + 增量修改分流：区分 local（局部重生成）、slot_change（全网重生成）、global'
    '（全网重生成）三种修改类型，大幅降低 LLM 调用成本；',
    'POI 存续巡检 + 自动替换：定时通过高德 API 逐条验证全库 POI 状态，失效景点自动排除并推荐替代，'
    '保障推荐可靠性；',
    '三级确定性评测体系：Micro（约束单元）/ Macro（query 全通过）/ Final Pass Rate 三级指标，'
    '24 项约束全部由确定性代码打分，无需人工评审，结果可复现。',
]
for inn in innovations:
    add_para('• ' + inn, size=10.5, space_after=1, first_line_indent=0.74)

# ── 补充说明 ──
add_heading_styled('补充说明')
add_para(
    '本项目已持续迭代 12 个大 Phase，从 14 天 MVP 演进为接近产品级的旅行规划系统。'
    '所有代码开源，评测数据完整可追溯。详细技术文档见项目 README.md 及 docs/ 目录。',
    size=10.5, first_line_indent=0.74
)

# ── Save ──
output_path = '智游伴_智能体说明文档.docx'
doc.save(output_path)
print(f'Done: {output_path}')
